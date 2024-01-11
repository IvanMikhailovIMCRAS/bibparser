import os
from typing import Optional, Tuple

import bibtexparser
from docx import Document
from pydantic import BaseModel

# pip install python-docx
doc = Document()


class Journal(BaseModel):
    issn: str
    name: str
    imp_fac: float
    Q: int


class Article(BaseModel):
    ID: str
    title: str
    author: str
    journal: str
    volume: int
    number: Optional[int] = None
    pages: str
    month: int
    year: int
    doi: str
    issn: str


journal_list = [
    Journal.model_validate(
        {
            "issn": "1422-0067",
            "name": "International Journal of Molecular Science",
            "imp_fac": 5.6,
            "Q": 1,
        }
    ),
    Journal.model_validate(
        {
            "issn": "0303-402X",
            "name": "Colloid and Polymer Science",
            "imp_fac": 2.434,
            "Q": 2,
        }
    ),
    Journal.model_validate(
        {"issn": "1525-7797", "name": "Biomacromolecules", "imp_fac": 6.979, "Q": 1}
    ),
    Journal.model_validate(
        {
            "issn": "1022-1336",
            "name": "Macromolecular Rapid Communications",
            "imp_fac": 5.006,
            "Q": 1,
        }
    ),
    Journal.model_validate(
        {
            "issn": "1027-4510",
            "name": "Journal of Surface Investigation",
            "imp_fac": 0.206,
            "Q": 3,
        }
    ),
    Journal.model_validate(
        {"issn": "2073-4360", "name": "Polymers", "imp_fac": 5.0, "Q": 1}
    ),
    Journal.model_validate(
        {"issn": "1744-6848", "name": "Soft Matter", "imp_fac": 4.046, "Q": 1}
    ),
    Journal.model_validate(
        {"issn": "2313-7673", "name": "Biomimetics", "imp_fac": 3.743, "Q": 2}
    ),
    Journal.model_validate(
        {"issn": "2158-3226", "name": "AIP Advances", "imp_fac": 1.697, "Q": 2}
    ),
    Journal.model_validate(
        {"issn": "2161-1653", "name": "ACS Macro Letters", "imp_fac": 7.015, "Q": 1}
    ),
]

journal_dict = {x.issn: x for x in journal_list}

journal_list.sort(key=lambda x: x.issn, reverse=False)


with open(
    os.path.join(os.path.dirname(os.path.realpath(__file__)), "lab24.bib")
) as bibtex_file:
    bib_database = bibtexparser.load(bibtex_file)

articles_list = []

for b in bib_database.entries:
    try:
        a = Article.model_validate(b)
        articles_list.append(a)
    except Exception:
        print(f"Error in {b['ID']}, \t doi: {b['doi']}")

articles_list.sort(key=lambda x: x.month, reverse=False)

author_dict = {
    "Borisov": ["Борисов", "Олег", "Владимирович"],
    "Zhulina": ["Жулина", "Екатерина", "Борисовна"],
    "Neelov": ["Неелов", "Игорь", "Михайлович"],
    "Simonova": ["Симонова", "Мария", "Александровна"],
    "Ivanov": ["Иванов", "Иван", "Владимирович"],
    "Mikhailov": ["Михайлов", "Иван", "Викторович"],
    "Lukiev": ["Лукиев", "Иван", "Васильевич"],
    "Salamatova": ["Попова", "Татьяна", "Олеговна"],
    "Popova": ["Попова", "Татьяна", "Олеговна"],
}

tbl = []

print("всего статей: ", len(articles_list))

for a in articles_list:
    row = []
    row.append(a.title)
    surname = ""
    name = ""
    patronymic = ""
    for fio in author_dict:
        if fio in a.author:
            surname += f"{author_dict[fio][0]} \n"
            name += f"{author_dict[fio][1]} \n"
            patronymic += f"{author_dict[fio][2]} \n"
    row.append(surname)
    row.append(name)
    row.append(patronymic)
    row.append(a.journal)
    row.append(f"{a.year}.{a.month} \n Т.{a.volume} вып.{a.number}")
    row.append(str(round(journal_dict[a.issn].imp_fac, 3)))
    row.append(str(journal_dict[a.issn].Q))
    row.append(a.issn)
    tbl.append(tuple(row))
items = tuple(tbl)


# добавляем таблицу с одной строкой
# для заполнения названий колонок
table = doc.add_table(1, len(items[0]))
# определяем стиль таблицы
# table.style = 'Light Shading Accent 1'
# Получаем строку с колонками из добавленной таблицы
head_cells = table.rows[0].cells
# добавляем названия колонок
# for i, item in enumerate(['Кол-во', 'ID', 'Описание']):
#     p = head_cells[i].paragraphs[0]
#     # название колонки
#     p.add_run(item).bold = True
#     # выравниваем посередине
#     p.alignment = WD_ALIGN_PARAGRAPH.CENTER
# добавляем данные к существующей таблице
for it in items:
    # добавляем строку с ячейками к объекту таблицы
    cells = table.add_row().cells
    for i, item in enumerate(it):
        # вставляем данные в ячейки
        cells[i].text = str(item)
        # если последняя ячейка
        # if i == 2:
        #     # изменим шрифт
        #     cells[i].paragraphs[0].runs[0].font.name = 'Arial'
doc.save("test.docx")
